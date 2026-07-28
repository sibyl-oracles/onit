'''
# Copyright 2025 Rowel Atienza. All rights reserved.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

Local Search Toolkit — search pipeline for in-house documents.

Inspired by Mistral's Search Toolkit (https://mistral.ai/news/search-toolkit/):
a composable pipeline that unifies ingestion (parse -> chunk -> embed/index)
and retrieval (BM25 sparse, dense embeddings, hybrid fusion) behind one
interface, running entirely on local infrastructure — no cloud dependency.

Supported formats: .pdf, .md, .txt, .csv, .docx, .xlsx

Retrieval methods:
- bm25   : Okapi BM25 sparse lexical retrieval (pure Python, no extra deps)
- dense  : cosine similarity over embeddings from any OpenAI-compatible
           /embeddings endpoint (set ONIT_EMBEDDING_HOST + ONIT_EMBEDDING_MODEL)
- hybrid : reciprocal rank fusion of bm25 and dense rankings (default;
           falls back to bm25 when no embedding endpoint is configured)

Optional dependencies:
    pip install python-docx openpyxl   # .docx and .xlsx ingestion
'''

import contextlib
import hashlib
import json
import math
import os
import re
import tempfile
import time
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Sequence, Tuple

import logging

logger = logging.getLogger(__name__)

INDEX_FILENAME = "index.json"
INDEX_VERSION = 1

SUPPORTED_EXTENSIONS = {
    '.pdf', '.md', '.markdown', '.txt', '.text', '.csv',
    '.docx', '.xlsx', '.xlsm',
}

CONTENT_HASH_CHARS = 16           # sha256 prefix kept as the document identity

MAX_FILE_SIZE = 20 * 1024 * 1024  # skip files larger than 20MB
MAX_SHEET_ROWS = 5000             # cap rows read per spreadsheet sheet
EMBED_BATCH_SIZE = 64
RRF_K = 60                        # reciprocal rank fusion constant

DEFAULT_CHUNK_SIZE = 1600         # characters per chunk
DEFAULT_CHUNK_OVERLAP = 200       # character overlap between chunks

NAME_LABEL_WEIGHT = 1.0           # weight of the filename/folder field in BM25
MAX_CHUNKS_PER_FILE = 2           # chunks one document may fill in a result page


def cap_per_file(ordered: List[Any], top_k: int,
                 file_of: Callable[[Any], str]) -> List[Any]:
    """Take the top_k best-ranked items, letting no one document supply more
    than MAX_CHUNKS_PER_FILE of them while other documents are still waiting.

    Ranking is per chunk, so a long document that matches the query in many
    places wins many of the slots the caller sees. A corpus README cataloguing
    every file mentions every topic, in a hundred chunks, and took nearly the
    whole page for queries the documents it catalogues answer directly — the
    document that actually held the answer was ranked, but below the fold.

    Items past a document's cap are not dropped, only deferred: if there are
    not enough distinct documents to fill top_k, they backfill in rank order,
    so a corpus where one file genuinely holds every match still returns a
    full page.
    """
    kept: List[Any] = []
    overflow: List[Any] = []
    counts: Dict[str, int] = {}
    for item in ordered:
        path = file_of(item)
        if counts.get(path, 0) < MAX_CHUNKS_PER_FILE:
            counts[path] = counts.get(path, 0) + 1
            kept.append(item)
            if len(kept) == top_k:
                return kept
        else:
            overflow.append(item)
    return kept + overflow[:top_k - len(kept)]


def file_content_hash(file_path: str, block_size: int = 1 << 20) -> str:
    """Hash of the file's bytes — an identity that does not depend on where
    the file lives. Two copies of the same document (the shared corpus file
    and a session's copy of it) hash alike, which is what lets search results
    from separate indexes be recognized as the same document."""
    digest = hashlib.sha256()
    with open(file_path, "rb") as f:
        for block in iter(lambda: f.read(block_size), b""):
            digest.update(block)
    return digest.hexdigest()[:CONTENT_HASH_CHARS]


# =============================================================================
# INGESTION: PARSERS
# =============================================================================

def parse_document(file_path: str) -> List[Tuple[str, str]]:
    """Parse a document into a list of (location, text) blocks.

    Location is a human-readable pointer within the file (e.g. "page 3",
    "sheet Sales"). Raises ValueError for unsupported or unreadable files.
    """
    ext = Path(file_path).suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}. "
                         f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")

    if ext == '.pdf':
        return _parse_pdf(file_path)
    if ext == '.docx':
        return _parse_docx(file_path)
    if ext in ('.xlsx', '.xlsm'):
        return _parse_xlsx(file_path)
    return _parse_text(file_path)


def _parse_pdf(file_path: str) -> List[Tuple[str, str]]:
    from pypdf import PdfReader
    reader = PdfReader(file_path)
    blocks = []
    for i, page in enumerate(reader.pages, 1):
        text = (page.extract_text() or "").strip()
        if text:
            blocks.append((f"page {i}", text))
    return blocks


def _parse_docx(file_path: str) -> List[Tuple[str, str]]:
    try:
        import docx
    except ImportError:
        raise ValueError("python-docx not installed. Run: pip install python-docx")
    document = docx.Document(file_path)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    blocks = []
    if paragraphs:
        blocks.append(("body", "\n".join(paragraphs)))
    for idx, table in enumerate(document.tables, 1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append("\t".join(cells))
        if rows:
            blocks.append((f"table {idx}", "\n".join(rows)))
    return blocks


def _parse_xlsx(file_path: str) -> List[Tuple[str, str]]:
    try:
        import openpyxl
    except ImportError:
        raise ValueError("openpyxl not installed. Run: pip install openpyxl")
    workbook = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    blocks = []
    try:
        for sheet in workbook.worksheets:
            rows = []
            for row_num, row in enumerate(sheet.iter_rows(values_only=True)):
                if row_num >= MAX_SHEET_ROWS:
                    break
                cells = ["" if c is None else str(c).strip() for c in row]
                if any(cells):
                    rows.append("\t".join(cells))
            if rows:
                blocks.append((f"sheet {sheet.title}", "\n".join(rows)))
    finally:
        workbook.close()
    return blocks


def _parse_text(file_path: str) -> List[Tuple[str, str]]:
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        content = f.read().strip()
    return [("", content)] if content else []


# =============================================================================
# INGESTION: CHUNKING
# =============================================================================

def chunk_text(text: str, chunk_size: int = DEFAULT_CHUNK_SIZE,
               overlap: int = DEFAULT_CHUNK_OVERLAP) -> List[str]:
    """Split text into chunks of ~chunk_size characters with overlap.

    Packs whole paragraphs when possible; paragraphs longer than chunk_size
    are hard-split with overlapping windows.
    """
    if len(text) <= chunk_size:
        return [text] if text.strip() else []

    overlap = min(overlap, chunk_size // 2)
    paragraphs = [p for p in re.split(r'\n\s*\n', text) if p.strip()]

    pieces = []
    for para in paragraphs:
        if len(para) <= chunk_size:
            pieces.append(para)
        else:
            step = chunk_size - overlap
            for start in range(0, len(para), step):
                piece = para[start:start + chunk_size]
                if piece.strip():
                    pieces.append(piece)
                if start + chunk_size >= len(para):
                    break

    chunks = []
    current = ""
    for piece in pieces:
        if current and len(current) + len(piece) + 2 > chunk_size:
            chunks.append(current)
            # Carry trailing overlap into the next chunk for continuity
            current = current[-overlap:] + "\n\n" if overlap else ""
        current += piece + "\n\n"
    if current.strip():
        chunks.append(current)

    return [c.strip() for c in chunks if c.strip()]


# =============================================================================
# RETRIEVAL: BM25 (SPARSE)
# =============================================================================

_TOKEN_RE = re.compile(r"[a-z0-9]+")


def _fold_plural(word: str) -> str:
    """Fold a common English plural onto its singular form.

    BM25 matches terms literally, so "scholarships" and "scholarship" are
    unrelated terms to it — a query in one number scores nothing against a
    document written in the other. Documents are titled in the singular
    ("...-scholarship-...") far more often than questions are asked in it,
    which is exactly where the mismatch costs the most.

    Deliberately not a full stemmer: the rules below are the plural forms, and
    nothing else. A wider stemmer trades recall for precision on a corpus this
    small, and mis-stems ("meeting" -> "meet") are harder to reason about than
    a missed match. What matters is that queries and documents are folded by
    the same function, so an imperfect fold still lines the two sides up.
    """
    if len(word) > 4 and word.endswith("ies"):        # policies -> policy
        return word[:-3] + "y"
    if len(word) > 4 and (word.endswith("ches") or word.endswith("shes")
                          or (word.endswith("es") and word[-3] in "sxz")):
        return word[:-2]                              # boxes -> box
    if (len(word) > 3 and word.endswith("s")
            and not word.endswith("ss")               # process, class
            and not word.endswith("us")):             # campus, syllabus
        return word[:-1]                              # scholarships -> scholarship
    return word


def tokenize(text: str) -> List[str]:
    return [_fold_plural(w) for w in _TOKEN_RE.findall(text.lower())]


def _name_label(file_path: str) -> str:
    """The filename and its folder as searchable words.

    Documents are often named for what they are about — a query naming the
    topic can otherwise score zero against every chunk of the file whose name
    matches it, and a zero-scoring chunk is dropped from the ranking entirely.
    The containing folder carries the same kind of signal one level up
    ("AI-Program/" for a query about the AI program).

    Only the immediate parent is used. Deeper ancestors are mostly generic
    ("docs", "Documents", a home directory) and would put the same tokens on
    every chunk in the index; the folder name is shared by its own files
    either way, so within a folder the filename stem still does the ranking.
    """
    path = Path(file_path)
    parts = [path.parent.name, path.stem]
    return " ".join(p.replace('-', ' ').replace('_', ' ') for p in parts if p)


def _labelled(chunk: Dict[str, Any]) -> str:
    """Chunk text with its folder and filename words prepended.

    Used for embedding only. BM25 scores the label as a separate field
    (see ``_bm25_ranking``) rather than reading it out of this string.
    """
    label = chunk.get("name_label", "")
    return f"{label} {chunk['text']}" if label else chunk["text"]


class BM25:
    """Okapi BM25 over a tokenized corpus. Pure Python, no dependencies."""

    def __init__(self, corpus_tokens: List[List[str]], k1: float = 1.5, b: float = 0.75):
        self.k1 = k1
        self.b = b
        self.corpus_tokens = corpus_tokens
        self.doc_lens = [len(tokens) for tokens in corpus_tokens]
        self.avg_len = (sum(self.doc_lens) / len(self.doc_lens)) if self.doc_lens else 0.0

        df: Dict[str, int] = {}
        for tokens in corpus_tokens:
            for term in set(tokens):
                df[term] = df.get(term, 0) + 1
        n = len(corpus_tokens)
        self.idf = {
            term: math.log(1 + (n - freq + 0.5) / (freq + 0.5))
            for term, freq in df.items()
        }

    def scores(self, query_tokens: List[str]) -> List[float]:
        results = [0.0] * len(self.corpus_tokens)
        for i, tokens in enumerate(self.corpus_tokens):
            if not tokens:
                continue
            tf: Dict[str, int] = {}
            for t in tokens:
                tf[t] = tf.get(t, 0) + 1
            norm = self.k1 * (1 - self.b + self.b * self.doc_lens[i] / (self.avg_len or 1))
            score = 0.0
            for term in query_tokens:
                if term not in tf:
                    continue
                freq = tf[term]
                score += self.idf.get(term, 0.0) * freq * (self.k1 + 1) / (freq + norm)
            results[i] = score
        return results


# =============================================================================
# RETRIEVAL: DENSE EMBEDDINGS (OPTIONAL)
# =============================================================================

def embedding_config() -> Optional[Dict[str, str]]:
    """Return embedding endpoint config from environment, or None if unset.

    Requires ONIT_EMBEDDING_HOST (OpenAI-compatible base URL) and
    ONIT_EMBEDDING_MODEL. ONIT_EMBEDDING_API_KEY is optional (local servers
    such as vLLM or Ollama usually don't need one).
    """
    host = os.environ.get('ONIT_EMBEDDING_HOST')
    model = os.environ.get('ONIT_EMBEDDING_MODEL')
    if not host or not model:
        return None
    key = (os.environ.get('ONIT_EMBEDDING_API_KEY')
           or os.environ.get('VLLM_API_KEY')
           or os.environ.get('OPENROUTER_API_KEY')
           or 'none')
    return {"host": host, "model": model, "key": key}


def embed_texts(texts: List[str]) -> Optional[List[List[float]]]:
    """Embed texts via the configured OpenAI-compatible endpoint.

    Returns None when no endpoint is configured or the request fails.
    """
    config = embedding_config()
    if not config or not texts:
        return None
    try:
        from openai import OpenAI
        client = OpenAI(base_url=config["host"], api_key=config["key"])
        vectors: List[List[float]] = []
        for start in range(0, len(texts), EMBED_BATCH_SIZE):
            batch = texts[start:start + EMBED_BATCH_SIZE]
            response = client.embeddings.create(model=config["model"], input=batch)
            vectors.extend(item.embedding for item in response.data)
        return vectors
    except Exception as e:
        logger.warning(f"Embedding request failed ({config['host']}): {e}")
        return None


def cosine_similarity(a: List[float], b: List[float]) -> float:
    dot = sum(x * y for x, y in zip(a, b))
    norm_a = math.sqrt(sum(x * x for x in a))
    norm_b = math.sqrt(sum(x * x for x in b))
    if norm_a == 0 or norm_b == 0:
        return 0.0
    return dot / (norm_a * norm_b)


def reciprocal_rank_fusion(rankings: List[List[int]], k: int = RRF_K) -> Dict[int, float]:
    """Fuse multiple rankings (lists of chunk indices, best first) into scores."""
    fused: Dict[int, float] = {}
    for ranking in rankings:
        for rank, idx in enumerate(ranking):
            fused[idx] = fused.get(idx, 0.0) + 1.0 / (k + rank + 1)
    return fused


# =============================================================================
# THE INDEX
# =============================================================================

class LocalSearchIndex:
    """Persistent chunk index over a corpus of local documents.

    The index lives as a single JSON file under index_dir. BM25 statistics
    are rebuilt in memory on load; embeddings (when available) are stored
    alongside each chunk.
    """

    def __init__(self, index_dir: str):
        self.index_dir = os.path.abspath(os.path.expanduser(index_dir))
        self.index_path = os.path.join(self.index_dir, INDEX_FILENAME)
        self.chunk_size = DEFAULT_CHUNK_SIZE
        self.chunk_overlap = DEFAULT_CHUNK_OVERLAP
        self.embedding_model: Optional[str] = None
        self.documents: Dict[str, Dict[str, Any]] = {}
        self.chunks: List[Dict[str, Any]] = []
        # Built together and invalidated together: _bm25 being None is what
        # signals these fields need rebuilding.
        self._bm25: Optional[BM25] = None
        self._bm25_label: Optional[BM25] = None
        self._leading: Optional[Dict[str, int]] = None
        self.load()

    # -- persistence ---------------------------------------------------------

    def load(self) -> None:
        if not os.path.isfile(self.index_path):
            return
        try:
            with open(self.index_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
        except Exception as e:
            logger.error(f"Failed to load index {self.index_path}: {e}")
            return
        # An index written by an older layout is not migrated: it is left
        # unread, so the next ingest rebuilds it from the documents instead of
        # serving chunks whose shape this version no longer understands.
        if data.get("version") != INDEX_VERSION:
            logger.warning(
                f"Ignoring index {self.index_path}: version {data.get('version')} "
                f"!= {INDEX_VERSION}; it will be rebuilt from the corpus")
            return
        self.chunk_size = data.get("chunk_size", DEFAULT_CHUNK_SIZE)
        self.chunk_overlap = data.get("chunk_overlap", DEFAULT_CHUNK_OVERLAP)
        self.embedding_model = data.get("embedding_model")
        self.documents = data.get("documents", {})
        self.chunks = data.get("chunks", [])
        # name_label was added after this index version shipped, so chunks
        # written by an earlier build carry none. Backfilling it here keeps
        # filename terms searchable without discarding embeddings to a rebuild.
        for chunk in self.chunks:
            if "name_label" not in chunk:
                chunk["name_label"] = _name_label(chunk["path"])
        self._bm25 = None
        self._leading = None

    def save(self) -> None:
        """Persist the index, atomically.

        The index is written to a temporary file in the same directory and
        renamed over the old one, so a concurrent reader — another onit
        service sharing this data directory, which loads the index at its own
        start — sees either the previous index or the new one, never the
        truncated middle of a write.
        """
        os.makedirs(self.index_dir, mode=0o700, exist_ok=True)
        data = {
            "version": INDEX_VERSION,
            "updated": time.strftime('%Y-%m-%d %H:%M:%S'),
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
            "embedding_model": self.embedding_model,
            "documents": self.documents,
            "chunks": self.chunks,
        }
        # mkstemp creates the file 0o600, and the leading dot keeps a partial
        # write out of any corpus walk over the directory holding the index.
        fd, tmp_path = tempfile.mkstemp(dir=self.index_dir, prefix=".index-",
                                        suffix=".tmp")
        try:
            with os.fdopen(fd, 'w', encoding='utf-8') as f:
                json.dump(data, f)
            os.replace(tmp_path, self.index_path)
        except Exception:
            with contextlib.suppress(OSError):
                os.unlink(tmp_path)
            raise

    # -- ingestion -----------------------------------------------------------

    def index_directory(self, directory: str, recursive: bool = True,
                        rebuild: bool = False,
                        exclude: Optional[Sequence[str]] = None) -> Dict[str, Any]:
        """Index all supported documents under directory. Skips unchanged
        files, drops entries for deleted files, and persists the result.

        ``exclude`` lists subtrees to leave alone — used when a directory
        nested under ``directory`` belongs to a different index, so the two do
        not both hold the same document.
        """
        directory = os.path.abspath(os.path.expanduser(directory))
        excluded = [os.path.abspath(os.path.expanduser(p)) for p in (exclude or [])]
        if rebuild:
            self.documents = {}
            self.chunks = []

        def _is_excluded(path: str) -> bool:
            return any(path == root or path.startswith(root + os.sep)
                       for root in excluded)

        pattern = "**/*" if recursive else "*"
        candidates = []
        for path in sorted(Path(directory).glob(pattern)):
            if not path.is_file():
                continue
            if any(part.startswith('.') for part in path.parts):
                continue
            if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
                continue
            if _is_excluded(str(path)):
                continue
            candidates.append(str(path))

        # Drop index entries for files that no longer exist under this root, and
        # for files an exclusion has since handed to another index — an earlier
        # pass may have ingested them while they were still in scope.
        removed = [p for p in self.documents
                   if p.startswith(directory + os.sep)
                   and (not os.path.isfile(p) or _is_excluded(p))]
        for p in removed:
            self._remove_document(p)

        indexed, skipped, errors = [], [], []
        for file_path in candidates:
            try:
                stat = os.stat(file_path)
                if stat.st_size > MAX_FILE_SIZE:
                    errors.append({"file": file_path, "error": "file too large (>20MB)"})
                    continue
                existing = self.documents.get(file_path)
                if existing and existing.get("mtime") == stat.st_mtime \
                        and existing.get("size") == stat.st_size:
                    skipped.append(file_path)
                    continue
                self._index_file(file_path, stat)
                indexed.append(file_path)
            except Exception as e:
                errors.append({"file": file_path, "error": str(e)})

        if indexed or removed:
            self._embed_pending()
            self._bm25 = None
            self._leading = None
            self.save()

        return {
            "directory": directory,
            "indexed": indexed,
            "skipped_unchanged": len(skipped),
            "removed": removed,
            "errors": errors,
            "no_text_extracted": self._no_text_extracted(),
            "total_documents": len(self.documents),
            "total_chunks": len(self.chunks),
            "embedding_model": self.embedding_model,
        }

    def _no_text_extracted(self) -> List[str]:
        """Documents that parsed without error but yielded no text.

        Overwhelmingly scanned PDFs: pages that are images carry no text layer
        for pypdf to extract, so the file is ingested, recorded, counted in
        total_documents — and matches nothing, because it contributed no chunk
        to search. That reads as a corpus with no answer rather than as a file
        needing OCR, so it is reported instead of passing silently.
        """
        return sorted(path for path, info in self.documents.items()
                      if not info.get("num_chunks"))

    def _index_file(self, file_path: str, stat: os.stat_result) -> None:
        blocks = parse_document(file_path)
        self._remove_document(file_path)
        name_label = _name_label(file_path)
        count = 0
        for location, text in blocks:
            for piece in chunk_text(text, self.chunk_size, self.chunk_overlap):
                self.chunks.append({
                    "path": file_path,
                    "location": location,
                    "name_label": name_label,
                    "text": piece,
                })
                count += 1
        self.documents[file_path] = {
            "mtime": stat.st_mtime,
            "size": stat.st_size,
            "hash": file_content_hash(file_path),
            "format": Path(file_path).suffix.lstrip('.').lower(),
            "num_chunks": count,
        }

    def _remove_document(self, file_path: str) -> None:
        self.documents.pop(file_path, None)
        self.chunks = [c for c in self.chunks if c["path"] != file_path]

    def _embed_pending(self) -> None:
        """Embed chunks that don't have vectors yet (no-op when no endpoint)."""
        config = embedding_config()
        if not config:
            return
        pending = [c for c in self.chunks if "embedding" not in c]
        if not pending:
            return
        vectors = embed_texts([_labelled(c) for c in pending])
        if vectors is None:
            return
        for chunk, vector in zip(pending, vectors):
            chunk["embedding"] = vector
        self.embedding_model = config["model"]

    # -- retrieval -----------------------------------------------------------

    def search(self, query: str, top_k: int = 5,
               method: str = "hybrid") -> List[Dict[str, Any]]:
        """Search the index. Returns ranked results with source metadata."""
        if not self.chunks:
            return []
        top_k = max(1, min(top_k, 20))

        bm25_ranking = self._bm25_ranking(query)
        dense_ranking = self._dense_ranking(query) if method in ("dense", "hybrid") else None

        if method == "dense" and dense_ranking is None:
            raise ValueError(
                "Dense retrieval unavailable: index has no embeddings or no "
                "embedding endpoint is configured (set ONIT_EMBEDDING_HOST "
                "and ONIT_EMBEDDING_MODEL)")

        if method == "bm25" or dense_ranking is None:
            ranked = bm25_ranking
            method_used = "bm25"
        elif method == "dense":
            ranked = dense_ranking
            method_used = "dense"
        else:
            fused = reciprocal_rank_fusion([
                [idx for idx, _ in bm25_ranking],
                [idx for idx, _ in dense_ranking],
            ])
            ranked = sorted(fused.items(), key=lambda item: -item[1])
            method_used = "hybrid"

        ranked = self._promote_document_openings(query, ranked)
        ranked = cap_per_file(ranked, top_k,
                              lambda item: self.chunks[item[0]]["path"])

        results = []
        for rank, (idx, score) in enumerate(ranked, 1):
            chunk = self.chunks[idx]
            results.append({
                "rank": rank,
                "score": round(float(score), 6),
                "file": chunk["path"],
                "location": chunk["location"],
                "text": chunk["text"],
                "method": method_used,
            })
        return results

    def _leading_chunk(self, path: str) -> Optional[int]:
        """Index of a document's first chunk — the one carrying its title.

        Chunks are appended in reading order at ingestion, so the first chunk
        recorded for a path is the opening of the document.
        """
        if self._leading is None:
            self._leading = {}
            for i, chunk in enumerate(self.chunks):
                self._leading.setdefault(chunk["path"], i)
        return self._leading.get(path)

    def _promote_document_openings(
            self, query: str,
            ranked: List[Tuple[int, float]]) -> List[Tuple[int, float]]:
        """Give each document matched by name a slot for its opening chunk.

        A document is retrieved as a whole but shown as an excerpt, and the
        excerpt is chosen by term frequency — which selects for repetition of
        the query words, not for the passage that identifies the document.
        A scholarship document whose name matched the query was returned twice,
        once on a passage listing degree programs and once on a page of contact
        links, while the chunks stating what the scholarship covers and who may
        apply ranked below its own cap and were never shown. Read from those two
        excerpts the document looks like unrelated boilerplate, so the answer
        was written from the web instead, omitting the one in-house document on
        the subject.

        The opening chunk is what makes a document recognizable: a title, and
        usually the summary under it. Promotion applies only where the corpus
        already voted for the document as a whole — its filename or folder
        matched the query — and the opening takes one of that document's own
        capped slots rather than a slot from another document, so a page still
        shows the same number of files. A document already led by its opening
        keeps the order it had.
        """
        if not ranked or self._bm25_label is None:
            return ranked
        label_scores = self._bm25_label.scores(tokenize(query))
        promoted: List[Tuple[int, float]] = []
        seen_paths: set = set()
        moved: set = set()
        for idx, score in ranked:
            if idx in moved:
                continue        # already emitted ahead of its document
            path = self.chunks[idx]["path"]
            if path not in seen_paths:
                seen_paths.add(path)
                opening = self._leading_chunk(path)
                if (opening is not None and opening != idx
                        and label_scores[idx] > 0):
                    # Ahead of the document's other chunks, so the per-file cap
                    # keeps it — ranked lower it was deferred as the document's
                    # third chunk and never shown. It stands for the document
                    # rather than matching on its own, so it inherits the rank
                    # the document earned.
                    promoted.append((opening, score))
                    moved.add(opening)
            promoted.append((idx, score))
        return promoted

    def _bm25_ranking(self, query: str) -> List[Tuple[int, float]]:
        """Rank chunks by BM25 over two fields: the passage text and the
        folder/filename label, summed with NAME_LABEL_WEIGHT.

        Scoring the label as its own field rather than as words prepended to
        the passage is what makes a document's name count. Concatenated, the
        label's handful of tokens sit in a chunk of several hundred: BM25
        saturates term frequency and divides by document length, so a title
        that names the query topic outright scored no better than the same word
        mentioned once in passing. A corpus index — a README cataloguing every
        filename in the corpus — beat the documents it catalogues on their own
        subjects, because it mentions every topic and is long enough to mention
        each one several times.

        As its own short field the label is scored on its own length, so
        matching two words of a four-word title is a strong signal, and it
        stays a bounded addition rather than an override: a passage that
        genuinely answers the query still outranks a file merely named for it.
        """
        if self._bm25 is None:
            self._bm25 = BM25([tokenize(c["text"]) for c in self.chunks])
            self._bm25_label = BM25(
                [tokenize(c.get("name_label", "")) for c in self.chunks])
        query_tokens = tokenize(query)
        text_scores = self._bm25.scores(query_tokens)
        label_scores = self._bm25_label.scores(query_tokens)
        ranking = [(i, t + NAME_LABEL_WEIGHT * l)
                   for i, (t, l) in enumerate(zip(text_scores, label_scores))
                   if t + NAME_LABEL_WEIGHT * l > 0]
        ranking.sort(key=lambda item: -item[1])
        return ranking

    def _dense_ranking(self, query: str) -> Optional[List[Tuple[int, float]]]:
        embedded = [(i, c["embedding"]) for i, c in enumerate(self.chunks)
                    if "embedding" in c]
        if not embedded:
            return None
        vectors = embed_texts([query])
        if not vectors:
            return None
        query_vector = vectors[0]
        ranking = [(i, cosine_similarity(query_vector, v)) for i, v in embedded]
        ranking.sort(key=lambda item: -item[1])
        return ranking

    # -- status --------------------------------------------------------------

    def status(self) -> Dict[str, Any]:
        by_format: Dict[str, int] = {}
        for info in self.documents.values():
            fmt = info.get("format", "unknown")
            by_format[fmt] = by_format.get(fmt, 0) + 1
        return {
            "index_path": self.index_path,
            "total_documents": len(self.documents),
            "total_chunks": len(self.chunks),
            "documents_by_format": by_format,
            "no_text_extracted": self._no_text_extracted(),
            "embedding_model": self.embedding_model,
            "embedded_chunks": sum(1 for c in self.chunks if "embedding" in c),
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
            "supported_formats": sorted(SUPPORTED_EXTENSIONS),
        }
